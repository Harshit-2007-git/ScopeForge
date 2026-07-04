"""
Word-document exporter — builds a formatted .docx from a ProjectScope object.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from schema import ProjectScope


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_docx(scope: ProjectScope) -> io.BytesIO:
    """Generate a Word document from the structured scope and return it as
    an in-memory bytes buffer ready for streaming."""

    doc = Document()

    # ── Styles ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ── Header ────────────────────────────────────────────────────────
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = (
        f"{scope.project_title}  |  "
        f"Generated {datetime.now(timezone.utc).strftime('%B %d, %Y')}"
    )
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.runs[0]
    header_run.font.size = Pt(9)
    header_run.font.italic = True

    # ── Title ─────────────────────────────────────────────────────────
    title = doc.add_heading(scope.project_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ── Summary ───────────────────────────────────────────────────────
    _add_heading(doc, "Summary", level=1)
    doc.add_paragraph(scope.summary)

    # ── Scope ─────────────────────────────────────────────────────────
    _add_heading(doc, "Scope", level=1)

    _add_heading(doc, "Included", level=2)
    _add_bullet_list(doc, scope.scope.included)

    _add_heading(doc, "Excluded", level=2)
    _add_bullet_list(doc, scope.scope.excluded)

    # ── Tech Stack ────────────────────────────────────────────────────
    if scope.tech_stack:
        _add_heading(doc, "Tech Stack", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light List Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Layer"
        hdr_cells[1].text = "Choice"
        hdr_cells[2].text = "Reason"
        for item in scope.tech_stack:
            row_cells = table.add_row().cells
            row_cells[0].text = item.layer
            row_cells[1].text = item.choice
            row_cells[2].text = item.reason

    # ── Timeline ──────────────────────────────────────────────────────
    if scope.timeline:
        _add_heading(doc, "Timeline", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light List Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Phase"
        hdr_cells[1].text = "Duration"
        hdr_cells[2].text = "Deliverable"
        for phase in scope.timeline:
            row_cells = table.add_row().cells
            row_cells[0].text = phase.phase
            row_cells[1].text = phase.duration
            row_cells[2].text = phase.deliverable

    # ── Risks & Questions ─────────────────────────────────────────────
    if scope.risks_and_questions:
        _add_heading(doc, "Risks & Open Questions", level=1)

        risks = [r for r in scope.risks_and_questions if r.type == "risk"]
        questions = [r for r in scope.risks_and_questions if r.type == "question"]

        if risks:
            _add_heading(doc, "Risks", level=2)
            _add_bullet_list(doc, [r.detail for r in risks])

        if questions:
            _add_heading(doc, "Open Questions", level=2)
            _add_bullet_list(doc, [q.detail for q in questions])

    # ── Save to buffer ────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
